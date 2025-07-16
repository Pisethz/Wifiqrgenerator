import subprocess
import re
import qrcode
import os
import openpyxl
from openpyxl import Workbook
from docx import Document
import webbrowser
import string

EXPORT_DIR = "wifi_exports"
os.makedirs(EXPORT_DIR, exist_ok=True)


# Helper to remove illegal Excel characters (control chars, etc.)
def sanitize_excel(s):
    if not s:
        return s
    # Remove ASCII control characters (0-31 except tab, LF, CR) and DEL (127)
    return ''.join(ch for ch in s if (ch in string.printable and ord(ch) not in {11, 12}) or ord(ch) >= 32)


def get_wifi_passwords():
    try:
        profiles = subprocess.check_output("netsh wlan show profiles", shell=True).decode(errors='ignore')
    except Exception as e:
        print("Error: You may need to run this script as administrator.")
        return []
    wifi_names = re.findall("All User Profile *: (.*)", profiles)
    if not wifi_names:
        print("No WiFi profiles found or insufficient permissions.")
        return []
    wifi_list = []
    for name in wifi_names:
        name = name.strip()
        try:
            result = subprocess.check_output(f'netsh wlan show profile name="{name}" key=clear', shell=True).decode(errors='ignore')
            password = re.search("Key Content *: (.*)", result)
            # Parse authentication type
            auth_match = re.search(r'Authentication *: (.+)', result)
            if auth_match:
                auth = auth_match.group(1).strip().upper()
                if 'WPA' in auth:
                    qr_type = 'WPA'
                elif 'WEP' in auth:
                    qr_type = 'WEP'
                else:
                    qr_type = 'nopass'
            else:
                qr_type = 'WPA'  # Default fallback
            if password:
                pwd = password.group(1).strip()
                ssid_escaped = escape_wifi_qr(name)
                pwd_escaped = escape_wifi_qr(pwd)
                if qr_type == 'nopass':
                    qr_data = f'WIFI:T:nopass;S:{ssid_escaped};;'
                else:
                    qr_data = f'WIFI:T:{qr_type};S:{ssid_escaped};P:{pwd_escaped};;'
                print(f"DEBUG: SSID: '{name}'")
                print(f"DEBUG: Password: '{pwd}'")
                print(f"DEBUG: QR String: '{qr_data}'")
                img = qrcode.make(qr_data)
                filename = f"wifi_qr_{name}.png"
                filename = re.sub(r'[^a-zA-Z0-9_-]', '_', filename)
                img_path = os.path.join(EXPORT_DIR, filename)
                img.save(img_path)
            else:
                pwd = None
                img_path = None
            wifi_list.append({
                'ssid': name,
                'password': pwd,
                'qr': img_path
            })
        except Exception as e:
            wifi_list.append({
                'ssid': name,
                'password': None,
                'qr': None
            })
    return wifi_list


def export_txt(wifi_list):
    txt_path = os.path.join(EXPORT_DIR, "wifi_passwords.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        for wifi in wifi_list:
            f.write(f"SSID: {wifi['ssid']}\n")
            f.write(f"Password: {wifi['password'] or '[None/Hidden]'}\n")
            f.write("-" * 30 + "\n")
    return txt_path

def export_excel(wifi_list):
    xlsx_path = os.path.join(EXPORT_DIR, "wifi_passwords.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.append(["SSID", "Password"])
    for wifi in wifi_list:
        ssid = sanitize_excel(wifi['ssid'])
        pwd = sanitize_excel(wifi['password'] or '[None/Hidden]')
        ws.append([ssid, pwd])
    wb.save(xlsx_path)
    return xlsx_path

def export_word(wifi_list):
    docx_path = os.path.join(EXPORT_DIR, "wifi_passwords.docx")
    doc = Document()
    doc.add_heading("WiFi Passwords", 0)
    for wifi in wifi_list:
        ssid = sanitize_excel(wifi['ssid'])
        pwd = sanitize_excel(wifi['password'] or '[None/Hidden]')
        doc.add_paragraph(f"SSID: {ssid}")
        doc.add_paragraph(f"Password: {pwd}")
        doc.add_paragraph("-" * 30)
    doc.save(docx_path)
    return docx_path
def escape_wifi_qr(s):
    if not s:
        return ''
    # Only escape \, ;, ,, :
    return s.replace('\\', '\\\\').replace(';', '\;').replace(',', '\,').replace(':', '\:')

def export_html(wifi_list):
    html_path = os.path.join(EXPORT_DIR, "wifi_passwords.html")
    modals_html = ""
    html = '''
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>WiFi Passwords & QR Codes</title>
<style>
body {
  font-family: 'Segoe UI', Arial, sans-serif;
  background: linear-gradient(135deg, #23283a 0%, #1e3c72 100%);
  color: #f3f3f3;
  margin: 0;
  padding: 0;
  min-height: 100vh;
}
.container {
  max-width: 900px;
  margin: 40px auto;
  background: rgba(35, 40, 58, 0.98);
  border-radius: 18px;
  box-shadow: 0 8px 32px #000a;
  padding: 40px 32px 32px 32px;
  position: relative;
}
h1 {
  text-align: center;
  color: #6cf;
  letter-spacing: 2px;
  margin-bottom: 18px;
}
h2 {
  text-align: center;
  color: #fff;
  letter-spacing: 2px;
  margin-bottom: 18px;
}
.search-bar {
  display: flex;
  justify-content: center;
  margin-bottom: 32px;
}
#searchInput {
  width: 340px;
  padding: 12px 18px;
  border-radius: 24px;
  border: none;
  font-size: 1.1em;
  background: #181c24;
  color: #fff;
  box-shadow: 0 2px 8px #0004;
  outline: none;
  transition: box-shadow 0.2s;
}
#searchInput:focus {
  box-shadow: 0 4px 16px #6cf6;
}
.export-bar {
  text-align: center;
  margin-bottom: 24px;
}
.export-bar a {
  display: inline-block;
  margin: 0 10px;
  background: #6cf;
  color: #222;
  padding: 10px 22px;
  border-radius: 8px;
  text-decoration: none;
  font-weight: bold;
  font-size: 1.05em;
  box-shadow: 0 2px 8px #0003;
  transition: background 0.2s, color 0.2s;
}
.export-bar a:hover {
  background: #4ad;
  color: #fff;
}
.wifi-list {
  display: flex;
  flex-wrap: wrap;
  gap: 28px;
  justify-content: center;
}
.wifi-card {
  background: #23283a;
  border-radius: 16px;
  box-shadow: 0 4px 16px #0007;
  padding: 24px 18px 18px 18px;
  width: 250px;
  text-align: center;
  position: relative;
  transition: transform 0.18s, box-shadow 0.18s;
  border: 1.5px solid #2e3c5a;
}
.wifi-card:hover {
  transform: translateY(-8px) scale(1.03);
  box-shadow: 0 8px 32px #6cf5;
  border-color: #6cf;
}
.wifi-card h2 {
  font-size: 1.18em;
  color: #6cf;
  margin: 0 0 10px 0;
  letter-spacing: 1px;
}
.wifi-card button {
  background: #6cf;
  color: #222;
  border: none;
  border-radius: 8px;
  padding: 8px 20px;
  margin: 10px 0 0 0;
  cursor: pointer;
  font-weight: bold;
  font-size: 1em;
  transition: background 0.2s, color 0.2s;
  box-shadow: 0 2px 8px #0003;
}
.wifi-card button:hover {
  background: #4ad;
  color: #fff;
}
.qr-img {
  width: 120px;
  height: 120px;
  max-width: 100%;
  margin: 12px auto 0 auto;
  display: block;
  background: #fff;
  border-radius: 10px;
  box-shadow: 0 2px 8px #0002;
}
.modal {
  display: none;
  position: fixed;
  z-index: 100;
  left: 0;
  top: 0;
  width: 100vw;
  height: 100vh;
  background: rgba(0,0,0,0.7);
  justify-content: center;
  align-items: center;
}
.modal.show {
  display: flex;
}
.modal-content {
  background: #23283a;
  margin: 0;
  padding: 36px 24px 28px 24px;
  border-radius: 20px;
  width: 95vw;
  max-width: 400px;
  text-align: center;
  color: #fff;
  box-shadow: 0 8px 32px #000a;
  position: relative;
  display: flex;
  flex-direction: column;
  align-items: center;
}
.modal-content .qr-img {
  width: 220px;
  height: 220px;
  max-width: 90vw;
  margin: 18px auto 0 auto;
  border-radius: 16px;
}
.close {
  color: #aaa;
  float: right;
  font-size: 32px;
  font-weight: bold;
  position: absolute;
  right: 24px;
  top: 12px;
  cursor: pointer;
  z-index: 101;
}
.close:hover {
  color: #fff;
}
@media (max-width: 600px) {
  .container { padding: 10px; }
  .wifi-card { width: 98vw; min-width: 0; }
  .modal-content { max-width: 98vw; padding: 18px 4vw 18px 4vw; }
  .modal-content .qr-img { width: 140px; height: 140px; }
}
</style>
</head>
<body>
<div class="container">
<h2>🔥💎© 2025 Pisethz x JackyJackyHunt!!. All Rights Reserved 💎🔥</h2>
<h1>WiFi Passwords & QR Codes</h1>
<div class="search-bar">
  <input type="text" id="searchInput" placeholder="Search WiFi SSID..." onkeyup="filterWiFi()">
</div>
<div class="export-bar">
<a href="wifi_passwords.txt" download>Export TXT</a>
<a href="wifi_passwords.xlsx" download>Export Excel</a>
<a href="wifi_passwords.docx" download>Export Word</a>
</div>
<div class="wifi-list" id="wifiList">
'''
    for i, wifi in enumerate(wifi_list):
        html += f'''<div class="wifi-card">
<h2 class="ssid">{wifi['ssid']}</h2>
<p>Password: <b>{wifi['password'] or '[None/Hidden]'}</b></p>
'''
        if wifi['qr']:
            qr_rel = os.path.basename(wifi['qr'])
            html += f'<img src="{qr_rel}" class="qr-img" alt="QR for {wifi["ssid"]}"><br>'
            html += f'<button onclick="showModal(\'modal{i}\')">Show QR</button>'
            # Collect modal HTML to output after container
            modals_html += f'''
<div id="modal{i}" class="modal">
  <div class="modal-content">
    <span class="close" onclick="closeModal('modal{i}')">&times;</span>
    <h2 style="color:#6cf;">{wifi['ssid']}</h2>
    <img src="{qr_rel}" class="qr-img" alt="QR for {wifi['ssid']}"><br>
    <p>Password: <b>{wifi['password']}</b></p>
  </div>
</div>
'''
        html += '</div>'
    html += f'''
</div>
</div>
{modals_html}
<script>
function showModal(id) {{
  var modal = document.getElementById(id);
  if (modal) {{
    modal.classList.add('show');
  }}
}}
function closeModal(id) {{
  var modal = document.getElementById(id);
  if (modal) {{
    modal.classList.remove('show');
  }}
}}
window.onclick = function(event) {{
  var modals = document.getElementsByClassName('modal');
  for (var i = 0; i < modals.length; i++) {{
    if (event.target == modals[i]) {{
      modals[i].classList.remove('show');
    }}
  }}
}}
function filterWiFi() {{
  var input = document.getElementById('searchInput');
  var filter = input.value.toLowerCase();
  var cards = document.getElementsByClassName('wifi-card');
  for (var i = 0; i < cards.length; i++) {{
    var ssid = cards[i].getElementsByClassName('ssid')[0];
    if (ssid.innerText.toLowerCase().indexOf(filter) > -1) {{
      cards[i].style.display = '';
    }} else {{
      cards[i].style.display = 'none';
    }}
  }}
}}
</script>
</body>
</html>
'''
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)
    return html_path

if __name__ == "__main__":
    wifi_list = get_wifi_passwords()
    if not wifi_list:
        exit()
    txt_path = export_txt(wifi_list)
    xlsx_path = export_excel(wifi_list)
    docx_path = export_word(wifi_list)
    html_path = export_html(wifi_list)
    print(f"\nExported to: {txt_path}, {xlsx_path}, {docx_path}, {html_path}")
    print(f"QR codes saved in: {EXPORT_DIR}")
    print("\nNote: Run this script as administrator for best results.\nRequires: pip install qrcode pillow openpyxl python-docx\n")
    # Optionally open HTML in browser
    try:
        webbrowser.open('file://' + os.path.abspath(html_path))
    except Exception:
        pass